# -*- coding: utf-8 -*-
import json
import os
import copy
import tkinter as tk
import win32api
import win32print
from docx import Document
import pyautogui
from PIL import ImageTk
import time
import pyperclip

pyautogui.PAUSE = 0.3

def kill_process(process_name='portal.exe'):
    '''
    结束进程portal.exe
    :param process_name: 默认为portal.exe
    :return:
    '''
    os.system(r'taskkill /F /IM %s' % process_name)


def run_process(process_path=r'D:\portal4.1\portal.exe'):
    '''
    相当于cmd运行：切换到d盘，cd到portal4.1文件夹，start D:\portal4.1\portal.exe
    顺便删除生成的临时空文件database.set，Portal.ini
    :param process_path:
    :return:
    '''

    dir_ = os.path.dirname(process_path)
    base_= os.path.basename(process_path)
    os.system("cd/d %s&&start %s" % (dir_, base_))
    if os.path.exists("database.set"):
        os.remove("database.set")
    if os.path.exists("Portal.ini"):
        os.remove("Portal.ini")

def recogniser(*args, area=None):
    '''
    识别图片，其中点转化为坐标，并返回；如果长时间识别识别则抛出异常Image_not_found
    :param args:
    :param area: 识别的范围，默认为None即识别全屏
    :return: 图片中点的坐标
    '''
    flag = True
    cord = None
    n = 0
    while flag:
        for im in args:

            box = pyautogui.locateOnScreen(im, region=area)
            if not box:
                n += 1
                time.sleep(1)
                if n > 15:
                    kill_process()
                    run_process()
            else:
                cord = pyautogui.center(box)
                flag = False
                break
    return cord


def copy_and_paste(zh_to_type):
    '''
    pyautogui的typewrite输入中文会有bug
    通过将中文字符串写入剪切板
    并配合热键ctrl+v实现输入中文字符功能
    '''

    pyperclip.copy(zh_to_type)

    pyautogui.hotkey('ctrl', 'v')


def change_drug_ward(ward_name):
    # 点击空白处获得焦点
    pyautogui.click(700, 300)
    # 关闭前端窗口（如果是最终窗口由于软件本身设计问题alt+f4无效）
    pyautogui.hotkey('alt', 'f4')
    # 识别菜单栏“系统维护”，获取其中点坐标并单击
    sys_set_cord = recogniser('.\\im_base\\ward_change.png', area=(1, 26, 389, 18))
    pyautogui.click(sys_set_cord)
    pyautogui.press('down')
    pyautogui.press('enter')

    recogniser('.\\im_base\\ward_rcg.png')

    pyautogui.press('tab')
    pyautogui.press('tab')
    # 输入需要切换到的药房名称
    copy_and_paste(ward_name)
    pyautogui.hotkey('alt', 'o')
    pyautogui.press('enter')
    # 识别菜单栏“业务处理”，获取其中点坐标并单击
    finish_cord = recogniser('.\\im_base\\work_station.png', area=(1, 26, 389, 18))
    pyautogui.click(finish_cord)
    pyautogui.press('down')
    pyautogui.press('enter')


# 将本次开药结果以json形式储存到record.db
def write_date(info_dict, **kwargs):
    if isinstance(info_dict, dict):
        info_dict['time'] = time.strftime('%Y %m %d-%H:%M:%S', time.localtime())
        for k, v in kwargs.items():
            info_dict[k] = v
        with open('record.db', 'a', encoding='utf-8') as f:
            f.write('\n')
            json.dump(info_dict, f, ensure_ascii=False, sort_keys=True)
            f.close()
    else:
        return

def to_word(text_header, text_body, rx_kind, flag=True):
    # 将获取的信息存为word文档以便打印
    # 根据rx_kind变量值选择使用儿科模板或成人模板
    if rx_kind:
        d = Document('.\\modalPd.mod')
    else:
        d = Document('.\\modalAdl.mod')

    # 获取word文档内的表格列表
    t = d.tables
    # 第一个表格为病人基本信息，包括姓名、性别、年龄
    t1 = t[0]
    # 如果是手动输入的文本信息
    if flag:
        # 姓名值位于t1.cell(0, 1)，即第一行第2列；
        # 性别值位于t1.cell(0, 3)，即第一行第4列；
        # 年龄值位于t1.cell(0, 5)，即第一行第6列；
        for i in range(1, len(t1.columns), 2):
            t1.cell(0, i).text = ''
            # 对应text_header列表的index为0,1,2
            run = t1.cell(0, i).paragraphs[0].add_run(text_header[(i - 1) // 2])
            run.font.name = '宋体'
            run.font.size = 240000
    # 如果是自动获取的信息，则前两项姓名、性别为图片，需分别插入
    else:
        t1.cell(0, 1).text = ''
        t1.cell(0, 3).text = ''
        run_name = t1.cell(0, 1).paragraphs[0].add_run()
        run_name.add_picture('.\\temp\\name_im.png')
        run_gender = t1.cell(0, 3).paragraphs[0].add_run()
        run_gender.add_picture('.\\temp\\gender_im.png')
        '''
        这里其实最后只剩下年龄信息，为了跟上面保持一致，写法稍复杂一些，也便于以后扩展内容。
        其实等同于：
        t1.cell(0, 5).text = ''
        run = t1.cell(0, 5).paragraphs[0].add_run(text_header[2])
        run.font.name = '宋体'
        run.font.size = 240000
        '''
        for i in range(5, len(t1.columns), 2):
            t1.cell(0, i).text = ''
            run = t1.cell(0, i).paragraphs[0].add_run(text_header[(i - 1) // 2])
            run.font.name = '宋体'
            run.font.size = 240000
    # 第二个表格为开药信息
    t2 = t[1]
    l = 0
    for row in range(0, len(t2.rows)):
        t2.cell(0, row).text = ''
    for j in range(0, len(text_body)):
        run1 = t2.cell(0, l).paragraphs[0].add_run(text_body[j])
        run1.font.name = '宋体'
        run1.font.size = 240000
        l += 1

    d.save('.\\temp.docx')


def print_file(filename):
    win32api.ShellExecute(
        0,
        'print',
        filename,
        '/d:"%s"' % win32print.GetDefaultPrinter(),
        '.',
        0)
    os.remove(filename)


# 定义药品用法
class DrugUse(object):
    def __init__(self, w, a):
        self.w = w
        self.a = a
        with open('conf.db', 'r', encoding='utf-8') as f:
            self.res = json.load(f)

    def output(self, drug_name):
        # 药品用法依据：体重、年龄、固定用法
        if self.res[drug_name]['basis'] == 'weight':
            real_val = self.w
        elif self.res[drug_name]['basis'] == 'age':
            real_val = self.a
        else:
            return ['{}*{}'.format(drug_name, self.res[drug_name]['amount']),
                    '              sig: {}{} {} {} '.format(
                            self.res[drug_name]['normal_dosage'],
                            self.res[drug_name]['unit'],
                            self.res[drug_name]['freq'],
                            self.res[drug_name]['route']
                        )
                    ]

        # 求根据剂量范围列表的index求用量列表的index
        dosage_idx = None
        for i in self.res[drug_name]['range_list']:
            if real_val in range(*i):
                dosage_idx = self.res[drug_name]['range_list'].index(i)
                break
        if dosage_idx is not None:
            # 如果dosage_idx有值，则根据index求具体用量
            dot = self.res[drug_name]['dosage_list'][dosage_idx]
        else:
            dot = self.res[drug_name]['normal_dosage']

        return ['{}*{}'.format(drug_name, self.res[drug_name]['amount']),
                '              sig: {}{} {} {} '.format(
                        dot,
                        self.res[drug_name]['unit'],
                        self.res[drug_name]['freq'],
                        self.res[drug_name]['route']
                    )
                ]


# 自定义药品用法用量的UI
class DrugUsageConf(tk.Toplevel):
    def __init__(self):
        super().__init__()
        self.title('配置药品用法')

        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()

        ww = 600
        wh = 750
        self.geometry(
            "%dx%d+%d+%d" %
            (ww, wh, (sw - ww) / 2, (sh - wh) / 2))
        self.after(1, lambda: self.focus_force())
        self.note = {
            "weight": "体重",
            "age": "年龄",
        }
        self.flag = True

        self.hint = tk.StringVar()
        self.hint.set('提示：请填写药品的用法用量等详细信息，点击“存储”按钮结束一种药品的信息填写并暂存；\n点击“完成并退出”按钮，生成配置文件并退出。')
        self.hint_label = tk.Label(self, textvariable=self.hint, font=('微软雅黑', 10))
        self.hint_label.pack()
        drug_frame = tk.Frame(self)
        drug_frame.pack(pady=15)

        tk.Label(drug_frame, text='药品名称：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.drug_name = tk.Entry(drug_frame, width=20, font=('微软雅黑', 10))
        self.drug_name.insert(0, '脑蛋白水解物口服液10mL')
        self.drug_name.focus_set()

        self.drug_name.pack(side=tk.LEFT)

        tk.Label(drug_frame, text='×总量：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.amount = tk.Entry(drug_frame, width=4, font=('微软雅黑', 10))
        self.amount.insert(0, '1盒')
        self.amount.pack(side=tk.LEFT)

        usage_frame = tk.Frame(self)
        usage_frame.pack(pady=15)

        tk.Label(usage_frame, text='频次', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.freq = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.freq.insert(0, 'Tid')
        self.freq.pack(side=tk.LEFT)

        tk.Label(usage_frame, text='用药途径', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.route = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.route.insert(0, 'po')
        self.route.pack(side=tk.LEFT)

        tk.Label(usage_frame, text='单位', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.unit = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.unit.insert(0, 'mL')
        self.unit.pack(side=tk.LEFT)


        tk.Label(self, text='如何计算用量：', font=('微软雅黑', 10)).pack()

        self.radio_value = tk.StringVar()
        self.radio_value.set('fixed')
        radio_frame = tk.Frame(self)
        radio_frame.pack()

        set_radio_fixed = tk.Radiobutton(radio_frame, text='固定用量', variable=self.radio_value, value='fixed',
                                         command=self.destroy_fix)

        set_radio_fixed.pack(side=tk.LEFT)

        set_radio_weight = tk.Radiobutton(radio_frame, text='根据体重范围', variable=self.radio_value, value='weight',
                                          command=self.destroy_add)

        set_radio_weight.pack(side=tk.LEFT)
        set_radio_age = tk.Radiobutton(radio_frame, text='根据年龄范围', variable=self.radio_value, value='age',
                                       command=self.destroy_add)

        set_radio_age.pack(side=tk.LEFT)
        # 储存点击添加按钮后自动生成的组件
        self.set_list = []
        # 每种药品的用法汇总
        self.usage_info = {}
        # 最终结果，以药品名字为键，self.usage_info为值的字典，最后以json格式储存
        self.res = {}
        # 每种药物的用量范围汇总
        self.usage_range = {}
        # 初始化时候显示‘固定用量’
        self.destroy_fix()

        btn_frame = tk.Frame(self)
        btn_frame.pack(pady=10)

        self.btn_ok = tk.Button(btn_frame, text='存储', width=12, height=2, activebackground='grey',
                           relief='groove')
        # 按下按钮就获得焦点，使add_frame失去焦点后获取值的功能能充分实现
        self.btn_ok.bind('<Button-1>', self.shift_focus)
        # 松开按钮才触发功能，防止最后填写的值（add_frame只有在失去焦点后才获取值）不能被正确获取。
        self.btn_ok.bind('<ButtonRelease-1>', self.confirm)
        self.btn_ok.pack(side=tk.LEFT)

        self.btn_cancel = tk.Button(btn_frame, text='完成并退出', width=12, height=2, activebackground='grey',
                               relief='groove')

        self.btn_cancel.bind('<Button-1>', self.shift_focus)

        self.btn_cancel.bind('<ButtonRelease-1>', self.cancel)
        self.btn_cancel.pack(side=tk.RIGHT)

        # 给整个Entry类组件绑定事件和回调函数
        self.bind_class("Entry", "<FocusIn>", self.on_select)

    def add(self):
        # 储存用法依据
        self.usage_info['basis'] = self.radio_value.get()
        self.usage_info["normal_dosage"] = ''

        # 定义赋值函数
        def aquire(event):
            # 三个格都有值才赋值
            if starts.get() and ends.get() and dots.get():
                # 以组件的实例为键，用法范围及该范围对应的用量数据为值
                self.usage_range[event.widget] = [int(starts.get()), int(ends.get()), int(dots.get())]


        add_frame = tk.Frame(self)
        # 新添加的组件失去焦点后触发取值函数 aquire
        add_frame.bind('<FocusOut>', aquire)
        add_frame.pack(pady=10)
        self.set_list.append(add_frame)
        tk.Label(add_frame, text='%s范围:' % self.note[self.radio_value.get()], font=('微软雅黑', 10)).pack(side=tk.LEFT)
        starts = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        starts.pack(side=tk.LEFT)
        tk.Label(add_frame, text='到', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        ends = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        ends.pack(side=tk.LEFT)
        tk.Label(add_frame, text='  用量:', font=('微软雅黑', 10)).pack(side=tk.LEFT, padx=5)
        dots = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        dots.pack(side=tk.LEFT)

        # 设定新增组件大于10个时新增按钮消失，避免无限制增加组件。
        if len(self.set_list) < 18:
            btn = tk.Button(self, text='增加%s范围' % self.note[self.radio_value.get()], width=12, height=2, activebackground='grey', relief='groove')
            btn.bind('<ButtonRelease-1>', self.add_destroy)

            btn.pack()
            self.set_list.append(btn)

    def destroy_it(self):
        # 摧毁所有新增组件功能，摧毁前情况已储存的数据。
        self.usage_info.clear()
        self.usage_range.clear()
        for each in self.set_list:
            each.destroy()
        self.set_list.clear()

    def destroy_add(self):
        self.destroy_it()
        self.add()

    def add_destroy(self, event):
        self.add()
        event.widget.destroy()

    def destroy_fix(self):
        # 点击固定用量时生成的组件
        self.destroy_it()
        self.usage_info['basis'] = self.radio_value.get()

        def aquire_fix(event):
            self.usage_info['normal_dosage'] = fix_dot.get()

        fix_frame = tk.Frame(self)
        fix_frame.bind('<FocusOut>', aquire_fix)
        fix_frame.pack(pady=10)
        self.set_list.append(fix_frame)
        tk.Label(fix_frame, text='固定用量：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        fix_dot = tk.Entry(fix_frame, font=('微软雅黑', 10), width=4)

        fix_dot.pack(side=tk.LEFT)

    def on_select(self, event):
        # event.widget 即触发事件的控件，可以直接用它的方法如.select_range()等
        event.widget.select_range(0, tk.END)

    def shift_focus(self, event):
        event.widget.focus_set()


    def confirm(self, event):

        self.usage_info["amount"] = self.amount.get()
        self.usage_info["freq"] = self.freq.get()
        self.usage_info["route"] = self.route.get()
        self.usage_info["unit"] = self.unit.get()
        self.usage_info.setdefault("normal_dosage", '')
        self.usage_info["range_list"] = []
        self.usage_info["dosage_list"] = []
        for each_list in self.usage_range.values():

            self.usage_info["range_list"].append(each_list[:2])
            self.usage_info["dosage_list"].append(each_list[-1])

        if not (self.usage_info["amount"] and self.usage_info["freq"] and self.usage_info["route"] and self.usage_info["unit"]):
            self.flag = False
        elif self.usage_info['basis'] == 'fixed':
            if not self.usage_info["normal_dosage"]:
                self.flag = False
            else:
                self.flag = True
        elif not (self.usage_info["range_list"] and self.usage_info["dosage_list"]):

                self.flag = False
        else:
            self.flag = True
        if self.flag:
            self.res[self.drug_name.get()] = copy.deepcopy(self.usage_info)

            self.radio_value.set('fixed')
            self.destroy_fix()
            self.hint_label.configure(fg='black')
            self.hint.set('药品【%s】用法已储存！\n请继续添加下一种药品信息；\n若需修改已填写药品信息，请保持药品名称不变，直接修改其他信息，后点击“储存！”\n完成全部药品的填写后请点击“完成并退出”按钮，生成配置文件并退出' % self.drug_name.get())

        else:
            self.hint_label.configure(fg='red')
            self.hint.set('错误：药品信息均不能为空，请修正！')

    def cancel(self, event):
        self.confirm(event)

        with open('conf.db', 'r', encoding='utf-8') as conf_file:
            conf_dict = json.load(conf_file)
        for i, v in self.res.items():
            conf_dict[i] = v
        conf_file.close()
        with open('conf.db', 'w+', encoding='utf-8') as conf_to_save:
            json.dump(conf_dict, conf_to_save, sort_keys=True, ensure_ascii=False)
        self.destroy()


class Ky(tk.Toplevel):

    # 注册关联药品名及用法
    with open('conf.db', 'r', encoding='utf-8') as f:
        res = json.load(f)
    drug = res.keys()

    def __init__(self):
        super().__init__()
        self.p_info = self.read_data()
        self.name, self.gender, self.age, self.weight = [
            self.p_info['name'],
            self.p_info['gender'],
            self.p_info['age'],
            self.p_info['weight']
        ]
        # 设定变量：true为手写的病人姓名性别；false为自动获取病人姓名性别
        self.flag = True
        self.text_to_w = []
        self.title('开具处方')

        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()

        ww = 650
        wh = 650
        self.geometry(
            "%dx%d+%d+%d" %
            (ww, wh, (sw - ww) / 2, (sh - wh) / 2))
        self.after(1, lambda: self.focus_force())
        self.rx_kind = False
        self.val = tk.StringVar()
        tk.Label(self, text='患者信息').grid(pady=5)
        self.f1 = tk.Frame(self)
        self.f1.grid(pady=5)

        self.pt_label = tk.Label(self.f1, text='姓    名')
        self.pt_label.grid(row=0, column=0, ipady=5)

        self.e1 = tk.Entry(self.f1)
        self.e1.insert(0, str(self.name))
        self.e1.bind('<FocusIn>', self.on_select)
        self.e1.focus_set()
        self.e1.grid(row=0, column=1)

        self.sex_label = tk.Label(self.f1, text='性    别')
        self.sex_label.grid(row=0, column=2, ipady=5)
        self.e2 = tk.Entry(self.f1)
        self.e2.insert(0, str(self.gender))
        self.e2.bind('<FocusIn>', self.on_select)
        self.e2.grid(row=0, column=3)
        tk.Label(self.f1, text='年龄(岁)').grid(row=1, column=0)
        self.e3 = tk.Entry(self.f1)
        self.e3.insert(0, str(self.age))
        self.e3.bind('<FocusIn>', self.on_select)
        self.e3.grid(row=1, column=1)
        tk.Label(self.f1, text='体重(kg)').grid(row=1, column=2)
        self.e4 = tk.Entry(self.f1)
        self.e4.insert(0, str(self.weight))
        self.e4.bind('<FocusIn>', self.on_select)
        self.e4.grid(row=1, column=3)
        self.button_frame = tk.Frame(self)
        self.button_frame.grid()

        self.bt_aquire = tk.Button(self.button_frame, text='获取患者信息', command=self.aquire_info, width=12,
                  height=1,
                  activebackground='grey',
                  relief='groove')
        self.bt_aquire.grid(row=0, column=0)

        self.bt_manual = tk.Button(self.button_frame, text='输入患者信息', command=self.manually, width=12,
                                   height=1,
                                   activebackground='grey',
                                   relief='groove')

        tk.Button(self.button_frame,
                  text='新增药品用法',
                  command=self.setup,
                  width=12,
                  height=1,
                  activebackground='grey',
                  relief='groove').grid(row=0, column=1)
        self.f = tk.Frame(self)
        self.f.grid(padx=40)
        self.croll = tk.Scrollbar(self.f, orient=tk.VERTICAL)
        self.croll.grid(row=0, column=1, sticky=tk.N + tk.S)
        self.val.set(tuple(sorted(self.drug)))
        self.lb = tk.Listbox(self.f,
                             listvariable=self.val,
                             selectmode=tk.MULTIPLE,
                             bd=1,
                             font=('微软雅黑', 10),
                             selectbackground='brown',
                             width=50,
                             height=9,
                             yscrollcommand=self.croll.set)
        self.lb.bind('<<ListboxSelect>>', self.show_info)
        self.lb.grid(row=0, column=0)
        self.croll.config(command=self.lb.yview)
        tk.Label(self, text='开药情况预览（可以直接修改预览信息后打印）').grid(pady=10)
        # text
        bg_frame_text = tk.Frame(self)
        bg_frame_text.grid(padx=30)
        scroll_bar_text = tk.Scrollbar(bg_frame_text, orient=tk.VERTICAL)
        scroll_bar_text.grid(row=0, column=1, sticky=tk.N + tk.S)
        self.text = tk.Text(bg_frame_text, width=80, height=20, yscrollcommand=scroll_bar_text.set)
        self.text.grid(row=0, column=0)
        scroll_bar_text.config(command=self.text.yview)

        self.fdrug = tk.Frame(self)
        self.fdrug.grid(pady=10)

        self.bt = tk.Button(self.fdrug,
                            text='打 印',
                            command=self.ok,
                            width=12,
                            height=2,
                            activebackground='grey',
                            relief='groove')
        self.bt.grid(row=0, column=0)
        self.btc = tk.Button(self.fdrug,
                             text='取 消',
                             command=self.can,
                             width=12,
                             height=2,
                             activebackground='grey',
                             relief='groove')
        self.btc.grid(row=0, column=1)

    def setup(self):
        '''
        弹出配置界面
        :return:
        '''
        # 实例化一个上层窗口
        setup_ui = DrugUsageConf()

        # 主窗口等待上层窗口结束，否则下面刷新文件列表的语句会马上执行。
        self.wait_window(setup_ui)
        # 配置结束后刷新配置文件列表
        with open('conf.db', 'r', encoding='utf-8') as f:
            res = json.load(f)
        self.val.set(tuple(sorted(res.keys())))

    def read_data(self):
        with open('record.db', encoding='utf-8') as f:
            lines = f.readlines()

            if lines:
                last_line = lines[-1]

                info_dict = json.loads(last_line)
                return info_dict
            else:
                return {"age": 0, "name": "", "weight": 0, "gender": "", "rx": []}

    def ok(self):

        self.age = int(self.e3.get())
        self.weight = int(self.e4.get())
        res = self.text.get(1.0, tk.END)
        self.text_to_w = res.split('\n')

        if self.age < 14:
            self.rx_kind = True
        else:
            self.rx_kind = False
        p_info = {}
        # 如果是手写输入的病人姓名性别
        if self.flag:
            self.name = self.e1.get()
            self.gender = self.e2.get()
            text_head = [self.name, self.gender, str(self.age)]
            write_date(
                p_info,
                name=self.name,
                age=self.age,
                weight=self.weight,
                gender=self.gender,
                rx=self.text_to_w
            )
        # 否则是自动截取的病人姓名性别
        else:
            text_head = [0, 0, str(self.age)]
            write_date(
                p_info,
                name='姓名',
                age=self.age,
                weight=self.weight,
                gender='性别',
                rx=self.text_to_w

            )


        text_body = self.text_to_w
        rx_kind = self.rx_kind
        to_word(text_head, text_body, rx_kind, flag=self.flag)

        print_file(os.path.abspath('.\\temp.docx'))
        self.destroy()

    def can(self):
        self.destroy()

    def on_select(self, event):
        # event.widget 即触发事件的控件，可以直接用它的方法如.select_range()等
        event.widget.select_range(0, tk.END)

    def aquire_info(self):
        # 自动获得病人姓名和性别
        # 隐藏原来的输入框
        self.e1.grid_forget()
        self.e2.grid_forget()
        # 截取屏幕上姓名和性别，并保存
        name_area = (244, 61, 61, 15)
        gender_area = (351, 61, 21, 15)
        name_im = pyautogui.screenshot(region=name_area)
        # 保存图片为了以后可以插入到word文件上
        name_im.save('.\\temp\\name_im.png')
        gender_im = pyautogui.screenshot(region=gender_area)
        gender_im.save('.\\temp\\gender_im.png')
        # 转换为tkinter认可的图片格式显示在原来的label控件上
        name_image = ImageTk.PhotoImage(name_im)
        gender_image = ImageTk.PhotoImage(gender_im)
        self.pt_label.config(image=name_image)
        self.pt_label.image = name_image
        self.sex_label.config(image=gender_image)
        self.sex_label.image = gender_image
        # 隐藏“获取病人信息”按钮
        self.bt_aquire.grid_forget()
        # 显示“输入病人信息”按钮
        self.bt_manual.grid(row=0, column=0)
        # 做标记
        self.flag = False

    def manually(self):
        # 手动输入病人姓名性别
        # 恢复输入框
        self.e1.grid(row=0, column=1)
        self.e2.grid(row=0, column=3)
        # 消除label标签上的图片，并更改文本显示
        self.pt_label.config(text='姓名', image='')
        self.sex_label.config(text='性别', image='')
        # 隐藏按钮
        self.bt_manual.grid_forget()
        # 恢复按钮
        self.bt_aquire.grid(row=0, column=0)
        # 做标记
        self.flag = True

    def show_info(self, event):
        try:
            self.name = self.e1.get()
            self.gender = self.e2.get()
            self.age = int(self.e3.get())
            self.weight = int(self.e4.get())
            drug_selected = []

            for each in self.lb.curselection():
                drug_selected.append(self.lb.get(each))

            text_to_show = []
            if drug_selected:

                drug_use = DrugUse(self.weight, self.age)
                for each in drug_selected:
                    each_usage = drug_use.output(each)

                    text_to_show.append('\n'.join(each_usage))

            text_data = '\n'.join(text_to_show)
            self.text.delete(1.0, 'end')
            self.text.insert('end', text_data)
        except IndexError:
            pass


class MainWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.wm_attributes('-topmost', 1)
        self.overrideredirect(True)
        sw = self.winfo_screenwidth()
        ww = 180
        wh = 20
        self.geometry(
            "%dx%d+%d+%d" %
            (ww, wh, sw - 260, 3))
        self.btn_frame = tk.Frame(self)
        self.btn_frame.pack()
        self.btn = tk.Button(self.btn_frame,
                             bg='white', relief='groove', text='手工处方', width=10, height=1,
                             activebackground='lightblue'
                             )
        self.btn.bind('<ButtonRelease-1>', self.start_app)
        self.btn.pack(side=tk.LEFT)

        self.btn_emergency = tk.Button(self.btn_frame,
                             bg='white', relief='groove', text='急诊', width=5, height=1,
                             activebackground='lightblue'
                             )
        self.btn_emergency.bind('<ButtonRelease-1>', self.emergency)
        self.btn_emergency.pack(side=tk.LEFT)

        self.btn_wm = tk.Button(self.btn_frame,
                                 bg='white', relief='groove', text='西药', width=5, height=1,
                                 activebackground='lightblue'
                                 )
        self.btn_wm.bind('<ButtonRelease-1>', self.wm)
        self.btn_wm.pack(side=tk.LEFT)


        self.btn_end = tk.Button(self.btn_frame,
                                 bg='white', relief='groove', text='退出', width=5, height=1,
                                 activebackground='red')
        self.btn_end.bind('<ButtonRelease-1>', self.end_app)
        self.btn_end.pack(side=tk.LEFT)

    def end_app(self, event):
        self.destroy()

    def start_app(self, event):
        start_ui = Ky()
        self.wait_window(start_ui)

    def emergency(self, event):
        change_drug_ward('急诊药房')

    def wm(self, event):
        change_drug_ward('门诊西药房')


if __name__ == '__main__':

    app = MainWindow()
    app.mainloop()


