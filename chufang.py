# -*- coding: utf-8 -*-
import json
import os, copy
import sys
import tkinter as tk
import win32api
import win32print

from docx import Document


def write_date(info_dict, name, age, weight, gender, rx_list):
    info_dict['name'] = name
    info_dict['age'] = age
    info_dict['weight'] = weight
    info_dict['rx'] = rx_list
    info_dict['gender'] = gender

    with open('record.db', 'a', encoding='utf-8') as f:
        f.write('\n')
        json.dump(info_dict, f, ensure_ascii=False, sort_keys=True)

        f.close()


# 定义药品用法
class DrugUse(object):
    def __init__(self, w, a):
        self.w = w
        self.a = a
        with open('conf.json', 'r', encoding='utf-8') as f:
            self.res = json.load(f)

    def output(self, drug_name):

        if self.res[drug_name]['basis'] == 'weight':
            real_val = self.w
        elif self.res[drug_name]['basis'] == 'age':
            real_val = self.a
        else:
            return ['{}*{}'.format(drug_name, self.res[drug_name]['amount']),
                    '              sig: {}{} {} {} '.format(
                            self.res[drug_name]['normal_dosage'],
                            self.res[drug_name]['unit'],
                            self.res[drug_name]['freq'],
                            self.res[drug_name]['route']
                        )
                    ]
        idx = None
        for i in self.res[drug_name]['range_list']:
            if real_val in range(*i):
                idx = self.res[drug_name]['range_list'].index(i)
        if idx:
            dot = self.res[drug_name]['dosage_list'][idx]
        else:
            dot = self.res[drug_name]['normal_dosage']

        return ['{}*{}'.format(drug_name, self.res[drug_name]['amount']),
                '              sig: {}{} {} {} '.format(
                        dot,
                        self.res[drug_name]['unit'],
                        self.res[drug_name]['freq'],
                        self.res[drug_name]['route']
                    )
                ]


class DrugUsageConf(tk.Toplevel):
    def __init__(self):
        super().__init__()
        self.title('配置药品用法')

        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()

        ww = 600
        wh = 700
        self.geometry(
            "%dx%d+%d+%d" %
            (ww, wh, (sw - ww) / 2, (sh - wh) / 2))
        self.after(1, lambda: self.focus_force())
        self.note = {
            "weight": "体重",
            "age": "年龄",
        }
        self.flag = True
        self.tool_entry = tk.Entry(self, font=('微软雅黑', 1), width=1)
        self.tool_entry.insert(0, '0')
        self.tool_entry.pack(side=tk.BOTTOM)

        self.hint = tk.StringVar()
        self.hint.set('提示：请填写药品的用法用量等详细信息，点击“存储”按钮结束一种药品的信息填写并暂存；\n点击“完成并退出”按钮，生成配置文件并退出。')
        tk.Label(self, textvariable=self.hint, font=('微软雅黑', 10)).pack()
        drug_frame = tk.Frame(self)
        drug_frame.pack(pady=15)

        tk.Label(drug_frame, text='药品名称：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.drug_name = tk.Entry(drug_frame, width=20, font=('微软雅黑', 10))
        self.drug_name.insert(0, '脑蛋白水解物口服液10mL')
        self.drug_name.focus_set()

        self.drug_name.pack(side=tk.LEFT)

        tk.Label(drug_frame, text='×总量：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.amount = tk.Entry(drug_frame, width=4, font=('微软雅黑', 10))
        self.amount.insert(0, '1盒')
        self.amount.pack(side=tk.LEFT)

        usage_frame = tk.Frame(self)
        usage_frame.pack(pady=15)

        tk.Label(usage_frame, text='频次', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.freq = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.freq.insert(0, 'Tid')
        self.freq.pack(side=tk.LEFT)

        tk.Label(usage_frame, text='用药途径', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.route = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.route.insert(0, 'po')
        self.route.pack(side=tk.LEFT)

        tk.Label(usage_frame, text='单位', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        self.unit = tk.Entry(usage_frame, width=4, font=('微软雅黑', 10))
        self.unit.insert(0, 'mL')
        self.unit.pack(side=tk.LEFT)



        tk.Label(self, text='如何计算用量：', font=('微软雅黑', 10)).pack()

        self.radio_value = tk.StringVar()
        self.radio_value.set('fixed')
        radio_frame = tk.Frame(self)
        radio_frame.pack()

        set_radio_fixed = tk.Radiobutton(radio_frame, text='固定用量', variable=self.radio_value, value='fixed',
                                         command=self.destroy_fix)

        set_radio_fixed.pack(side=tk.LEFT)

        set_radio_weight = tk.Radiobutton(radio_frame, text='根据体重范围', variable=self.radio_value, value='weight',
                                          command=self.destroy_add)

        set_radio_weight.pack(side=tk.LEFT)
        set_radio_age = tk.Radiobutton(radio_frame, text='根据年龄范围', variable=self.radio_value, value='age',
                                       command=self.destroy_add)

        set_radio_age.pack(side=tk.LEFT)
        self.set_list = []
        self.usage_info = {}
        self.res = {}
        self.usage_range = {}
        self.destroy_fix()
        btn_frame = tk.Frame(self)
        btn_frame.pack(pady=10)
        self.btn_ok = tk.Button(btn_frame, text='存储', width=12, height=2, activebackground='grey',
                           relief='groove')
        # 按下按钮就获得焦点，使add_frame失去焦点后获取值的功能能充分实现
        self.btn_ok.bind('<Button-1>', self.shift_focus)
        # 松开按钮才触发功能，防止最后填写的值（add_frame只有在失去焦点后才获取值）不能被正确获取。
        self.btn_ok.bind('<ButtonRelease-1>', self.confirm)
        self.btn_ok.pack(side=tk.LEFT)
        self.btn_cancel = tk.Button(btn_frame, text='完成并退出', width=12, height=2, activebackground='grey',
                               relief='groove')

        self.btn_cancel.bind('<Button-1>', self.shift_focus)

        self.btn_cancel.bind('<ButtonRelease-1>', self.cancel)
        self.btn_cancel.pack(side=tk.RIGHT)
        self.warn_text = tk.StringVar()
        tk.Label(self, textvariable=self.warn_text, font=('微软雅黑', 10), fg='red', width=25).pack()
        self.bind_class("Entry", "<FocusIn>", self.on_select)



    def add(self):
        self.usage_info['basis'] = self.radio_value.get()

        def aquire(event):
            # 三个格都有值才赋值
            if starts.get() and ends.get() and dots.get():
                self.usage_range[event.widget] = [starts.get(), ends.get(), dots.get()]
        add_frame = tk.Frame(self)
        add_frame.bind('<FocusOut>', aquire)
        add_frame.pack(pady=10)
        self.set_list.append(add_frame)
        tk.Label(add_frame, text='%s范围:' % self.note[self.radio_value.get()], font=('微软雅黑', 10)).pack(side=tk.LEFT)
        starts = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        starts.pack(side=tk.LEFT)
        tk.Label(add_frame, text='到', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        ends = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        ends.pack(side=tk.LEFT)
        tk.Label(add_frame, text='  用量:', font=('微软雅黑', 10)).pack(side=tk.LEFT, padx=5)
        dots = tk.Entry(add_frame, font=('微软雅黑', 10), width=4)

        dots.pack(side=tk.LEFT)
        if len(self.set_list) < 18:
            btn = tk.Button(self, text='增加%s范围' % self.note[self.radio_value.get()], width=12, height=2, activebackground='grey', relief='groove')
            btn.bind('<ButtonRelease-1>', self.add_destroy)

            btn.pack()
            self.set_list.append(btn)



    def destroy_it(self):

        self.usage_info.clear()
        self.usage_range.clear()
        for each in self.set_list:
            each.destroy()
        self.set_list.clear()

    def destroy_add(self):
        self.destroy_it()
        self.add()

    def add_destroy(self, event):
        self.add()
        event.widget.destroy()

    def destroy_fix(self):

        self.destroy_it()
        self.usage_info['basis'] = self.radio_value.get()

        def aquire_fix(event):
            self.usage_info['normal_dosage'] = fix_dot.get()

        fix_frame = tk.Frame(self)
        fix_frame.bind('<FocusOut>', aquire_fix)
        fix_frame.pack(pady=10)
        self.set_list.append(fix_frame)
        tk.Label(fix_frame, text='固定用量：', font=('微软雅黑', 10)).pack(side=tk.LEFT)
        fix_dot = tk.Entry(fix_frame, font=('微软雅黑', 10), width=4)

        fix_dot.pack(side=tk.LEFT)

    def on_select(self, event):
        # event.widget 即触发事件的控件，可以直接用它的方法如.select_range()等
        event.widget.select_range(0, tk.END)

    def shift_focus(self, event):
        event.widget.focus_set()


    def confirm(self, event):

        self.usage_info["amount"] = self.amount.get()
        self.usage_info["freq"] = self.freq.get()
        self.usage_info["route"] = self.route.get()
        self.usage_info["unit"] = self.unit.get()
        self.usage_info["normal_dosage"] = ''
        self.usage_info["range_list"] = []
        self.usage_info["dosage_list"] = []
        for each_list in self.usage_range.values():

            self.usage_info["range_list"].append(each_list[:2])
            self.usage_info["dosage_list"].append(each_list[-1])

        if not (self.usage_info["amount"] and self.usage_info["freq"] and self.usage_info["route"] and self.usage_info["unit"]):
            self.flag = False
        elif self.usage_info['basis'] == 'fixed':
            if not self.usage_info["normal_dosage"]:
                self.flag = False
        elif not (self.usage_info["range_list"] and self.usage_info["dosage_list"]):
                self.flag = False
        else:
            self.flag = True
        if self.flag:
            self.res[self.drug_name.get()] = copy.deepcopy(self.usage_info)

            self.radio_value.set('fixed')
            self.destroy_fix()
            self.hint.set('药品【%s】用法已储存！\n请继续添加下一种药品信息；\n若需修改已填写药品信息，请保持药品名称不变，直接修改其他信息，后点击“储存！”\n完成全部药品的填写后请点击“完成并退出”按钮，生成配置文件并退出' % self.drug_name.get())

        else:
            self.hint.set('错误：药品信息均不能为空，请修正！')

    def cancel(self, event):
        self.usage_info["amount"] = self.amount.get()
        self.usage_info["freq"] = self.freq.get()
        self.usage_info["route"] = self.route.get()
        self.usage_info["unit"] = self.unit.get()
        self.usage_info["normal_dosage"] = ''
        self.usage_info["range_list"] = []
        self.usage_info["dosage_list"] = []
        for each_list in self.usage_range.values():
            self.usage_info["range_list"].append(each_list[:2])
            self.usage_info["dosage_list"].append(each_list[-1])

        if not (self.usage_info["amount"] and self.usage_info["freq"] and self.usage_info["route"] and self.usage_info["unit"]):
            self.flag = False
        elif self.usage_info['basis'] == 'fixed':
            if not self.usage_info["normal_dosage"]:
                self.flag = False
        elif not (self.usage_info["range_list"] and self.usage_info["dosage_list"]):
            self.flag = False
        else:
            self.flag = True
        if self.flag:
            self.res[self.drug_name.get()] = copy.deepcopy(self.usage_info)

        with open('conf.db', 'r', encoding='utf-8') as conf_file:
            conf_dict = json.load(conf_file)
        for i, v in self.res.items():
            conf_dict[i] = v
        conf_file.close()
        with open('conf.db', 'w+', encoding='utf-8') as conf_to_save:
            json.dump(conf_dict, conf_to_save, sort_keys=True, ensure_ascii=False)
        self.destroy()
        sys.exit(0)


class Ky(object):
    # 注册关联药品名及用法
    with open('conf.db', 'r', encoding='utf-8') as f:
        res = json.load(f)
    drug = res.keys()

    def __init__(self, root):
        self.root = root
        self.p_info = self.read_data()
        self.name, self.gender, self.age, self.weight = [
            self.p_info['name'],
            self.p_info['gender'],
            self.p_info['age'],
            self.p_info['weight']
        ]

        self.text_to_w = []
        self.root.title('开具处方')
        # self.root.wm_attributes('-topmost', 1)
        # self.root.overrideredirect(True)

        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()

        ww = 650
        wh = 650
        self.root.geometry(
            "%dx%d+%d+%d" %
            (ww, wh, (sw - ww) / 2, (sh - wh) / 2))
        self.root.after(1, lambda: self.root.focus_force())
        self.val = tk.StringVar()
        tk.Label(self.root, text='患者信息').grid(pady=5)
        self.f1 = tk.Frame(self.root)
        self.f1.grid(pady=5)

        tk.Label(self.f1, text='姓    名').grid(row=0, column=0)
        self.e1 = tk.Entry(self.f1)
        self.e1.insert(0, str(self.name))
        self.e1.bind('<FocusIn>', self.on_select)
        self.e1.focus_set()
        self.e1.grid(row=0, column=1)
        tk.Label(self.f1, text='性    别').grid(row=0, column=2)
        self.e2 = tk.Entry(self.f1)
        self.e2.insert(0, str(self.gender))
        self.e2.bind('<FocusIn>', self.on_select)
        self.e2.grid(row=0, column=3)
        tk.Label(self.f1, text='年龄(岁)').grid(row=1, column=0)
        self.e3 = tk.Entry(self.f1)
        self.e3.insert(0, str(self.age))
        self.e3.bind('<FocusIn>', self.on_select)
        self.e3.grid(row=1, column=1)
        tk.Label(self.f1, text='体重(kg)').grid(row=1, column=2)
        self.e4 = tk.Entry(self.f1)
        self.e4.insert(0, str(self.weight))
        self.e4.bind('<FocusIn>', self.on_select)
        self.e4.grid(row=1, column=3)
        tk.Button(self.root, text='新增药品用法', command=self.setup, width=12,
                             height=1,
                             activebackground='grey',
                             relief='groove').grid()
        self.f = tk.Frame(self.root)
        self.f.grid(padx=40)
        self.croll = tk.Scrollbar(self.f, orient=tk.VERTICAL)
        self.croll.grid(row=0, column=1, sticky=tk.N + tk.S)
        self.val.set(tuple(sorted(self.drug)))
        self.lb = tk.Listbox(self.f,
                             listvariable=self.val,
                             selectmode=tk.MULTIPLE,
                             bd=1,
                             font=('微软雅黑', 10),
                             selectbackground='brown',
                             width=50,
                             height=9,
                             yscrollcommand=self.croll.set)
        self.lb.bind('<<ListboxSelect>>', self.show_info)
        self.lb.grid(row=0, column=0)
        self.croll.config(command=self.lb.yview)
        tk.Label(root, text='开药情况预览（可以直接修改预览信息后打印）').grid(pady=10)
        # text
        bg_frame_text = tk.Frame(root)
        bg_frame_text.grid(padx=30)
        scroll_bar_text = tk.Scrollbar(bg_frame_text, orient=tk.VERTICAL)
        scroll_bar_text.grid(row=0, column=1, sticky=tk.N + tk.S)
        self.text = tk.Text(bg_frame_text, width=80, height=20, yscrollcommand=scroll_bar_text.set)
        self.text.grid(row=0, column=0)
        scroll_bar_text.config(command=self.text.yview)

        self.fdrug = tk.Frame(self.root)
        self.fdrug.grid(pady=10)

        self.bt = tk.Button(self.fdrug,
                            text='打 印',
                            command=self.ok,
                            width=12,
                            height=2,
                            activebackground='grey',
                            relief='groove')
        self.bt.grid(row=0, column=0)
        self.btc = tk.Button(self.fdrug,
                             text='取 消',
                             command=self.can,
                             width=12,
                             height=2,
                             activebackground='grey',
                             relief='groove')
        self.btc.grid(row=0, column=1)


    def setup(self):
        '''
        弹出配置界面
        :return:
        '''
        # 实例化一个上层窗口
        setup_ui = DrugUsageConf()

        # 主窗口等待上层窗口结束，否则下面刷新文件列表的语句会马上执行。
        self.root.wait_window(setup_ui)
        # 配置结束后刷新配置文件列表
        with open('conf.db', 'r', encoding='utf-8') as f:
            res = json.load(f)
        self.drug = res.keys()


    def read_data(self):
        with open('record.db', encoding='utf-8') as f:
            lines = f.readlines()

            if lines:
                last_line = lines[-1]

                info_dict = json.loads(last_line)
                return info_dict
            else:
                return {"age": 0, "name": "", "weight": 0, "gender": "", "rx": []}

    def ok(self):
        self.name = self.e1.get()
        self.gender = self.e2.get()
        self.age = int(self.e3.get())
        self.weight = int(self.e4.get())
        res = self.text.get(1.0, tk.END)
        self.text_to_w = res.split('\n')

        self.root.destroy()

    def can(self):
        self.root.destroy()
        sys.exit(0)

    def on_select(self, event):
        # event.widget 即触发事件的控件，可以直接用它的方法如.select_range()等
        event.widget.select_range(0, tk.END)

    def show_info(self, event):
        try:
            self.name = self.e1.get()
            self.gender = self.e2.get()
            self.age = int(self.e3.get())
            self.weight = int(self.e4.get())
            drug_selected = []

            for each in self.lb.curselection():
                drug_selected.append(self.lb.get(each))

            text_to_show = []
            if drug_selected:

                drug_use = DrugUse(self.weight, self.age)
                for each in drug_selected:
                    each_usage = drug_use.output(each)

                    text_to_show.append('\n'.join(each_usage))

            text_data = '\n'.join(text_to_show)
            self.text.delete(1.0, 'end')
            self.text.insert('end', text_data)
        except IndexError:
            pass


def to_word(text_header, text_body):
    d = Document('.\\modal.mod')

    t = d.tables

    t1 = t[0]
    for i in range(1, len(t1.columns), 2):
        t1.cell(0, i).text = ''
        run = t1.cell(0, i).paragraphs[0].add_run(text_header[(i - 1) // 2])
        run.font.name = '宋体'
        run.font.size = 240000

    t2 = t[1]
    l = 0
    for row in range(0, len(t2.rows)):
        t2.cell(0, row).text = ''
    for j in range(0, len(text_body)):
        run1 = t2.cell(0, l).paragraphs[0].add_run(text_body[j])
        run1.font.name = '宋体'
        run1.font.size = 240000
        l += 1

    d.save('.\\temp.docx')


def print_file(filename):
    win32api.ShellExecute(
        0,
        'print',
        filename,
        '/d:"%s"' % win32print.GetDefaultPrinter(),
        '.',
        0)
    os.remove(filename)


if __name__ == '__main__':
    while True:
        root = tk.Tk()
        app = Ky(root)
        root.mainloop()

        p_info = {}
        write_date(p_info, app.name, app.age, app.weight, app.gender, app.text_to_w)

        text_head = [app.name, app.gender, str(app.age)]
        text_body = app.text_to_w
        to_word(text_head, text_body)

        print_file(os.path.abspath('.\\temp.docx'))
